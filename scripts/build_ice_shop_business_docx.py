from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "La_Neige_法式雪冰室創業計劃書.docx"
IMG = ROOT / "ice_shop_design" / "website" / "assets"


def set_font(run, size=11, bold=False, color="2B2420"):
    run.font.name = "Arial Unicode MS"
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Arial Unicode MS")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    set_font(r, 18 if level == 1 else 14, True, "8B1E3F" if level == 1 else "1F5C4B")
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, 10.5)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        set_font(r, 9.5, True, "FFFFFF")
        table.rows[0].cells[i]._tc.get_or_add_tcPr()
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            r = cells[i].paragraphs[0].add_run(str(value))
            set_font(r, 9)
    doc.add_paragraph()


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("La Neige 法式雪冰室創業計劃書")
set_font(r, 24, True, "8B1E3F")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("台灣冰品 x 法式生活感 x 父子退休創業教學案例")
set_font(r, 12, False, "1F5C4B")
doc.add_picture(str(IMG / "hero-storefront.png"), width=Inches(6.5))

add_heading(doc, "一、計畫摘要")
add_para(doc, "La Neige 法式雪冰室是一間結合台灣冰品、法式生活美感、父子創業故事與社群互動體驗的小型冰店。品牌由從英國學建築與藝術回台的琰綱，與退休後展開第二人生的父親共同打造。")
add_para(doc, "核心標語：把台灣的冰，做成一場法式午後。")

add_heading(doc, "二、品牌故事")
add_para(doc, "琰綱從英國帶回建築與藝術的視角，父親則把退休後對登山與重機的熱情帶進品牌。這間店不只賣冰，也把生活方式、家人關係與地方感轉化為可被分享的故事。")
doc.add_picture(str(IMG / "owner-father-normalized.jpg"), width=Inches(2.8))

add_heading(doc, "三、商品與冬季補強")
add_table(doc, ["分區", "商品方向", "價格"], [
    ["雪花冰", "焦糖伯爵、草莓玫瑰、芒果香草", "130-180"],
    ["刨冰", "檸檬塔刨冰、紅豆牛奶、芋泥奶霜", "100-150"],
    ["冬季飲品", "黑糖薑奶、桂圓紅棗茶、燒仙草奶茶、花生湯圓可可", "90-130"],
    ["湯品", "法式南瓜濃湯、番茄蔬菜湯", "90-120"],
])

add_heading(doc, "四、空間與設備")
add_para(doc, "租用 4 x 10 公尺空間，租金約 2 萬元。座位區規劃 2-4 人桌共 8 桌，並設透明製作區、法式拍照牆、直式電子海報輪播。設備包含商用製冰機、雪花冰機、刨冰機、冷凍櫃、冷藏櫃、淨水設備、POS 與行動支付。")

add_heading(doc, "五、網站與社群行銷")
add_para(doc, "網站範例包含品牌故事、菜單、拍照牆上傳、顧客投票留言、會員訂閱、阿山哥俱樂部。社群分工為：FB 做在地活動與優惠，IG 做美感與打卡，TikTok/抖音做短影音製作過程。")
add_table(doc, ["月份", "主題", "主打商品", "社群任務"], [
    ["2月", "玫瑰情人節", "草莓玫瑰雙人冰", "雙人合照、粉色 Reels"],
    ["4月", "伯爵茶香季", "焦糖伯爵雪花冰", "茶湯製作、午後茶冰"],
    ["7月", "芒果左岸", "芒果香草雪花冰", "芒果切片、限量倒數"],
    ["12月", "聖誕雪花季", "聖誕草莓雪花冰", "聖誕拍照牆、杯套小物"],
])

add_heading(doc, "六、會員與阿山哥俱樂部")
add_para(doc, "阿山哥俱樂部延伸父親登山與重機喜好。會員每月可 PO 爬山照，或 PO 騎士重機照並介紹機車品牌與車款，即可領取優惠碼。登山任務優惠碼 ASHAN月份，重機任務優惠碼 RIDER月份。")
doc.add_picture(str(IMG / "father-motorcycle.jpg"), width=Inches(4.2))

add_heading(doc, "七、財務試算")
add_table(doc, ["情境", "日客數", "客單", "月營收"], [
    ["保守", "40", "130", "約 15.6 萬"],
    ["穩定", "70", "130", "約 27.3 萬"],
    ["旺季", "100", "130", "約 39 萬"],
])
add_para(doc, "經營目標為食材成本控制在 30-38%，租金控制在營收 10% 以內，夏季冰品創造獲利，冬季以熱飲、湯品、會員活動補淡季。")

add_heading(doc, "八、GitHub 教學連結")
add_para(doc, "網站範例可放置於 GitHub Pages 的 docs/ 目錄。推上 GitHub 後，可在 Repository Settings > Pages 選擇 main branch / docs folder，即可取得公開教學網址。")

doc.save(OUT)
print(OUT)
