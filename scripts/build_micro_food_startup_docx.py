from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "micro_food_startup_teaching_plan.md"
OUTPUT = ROOT / "微型餐飲創業教學計劃_葉佳山.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "606A78"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "C9D3DF"


class NumberingManager:
    def __init__(self, doc):
        self.doc = doc
        self.numbering = doc.part.numbering_part.element
        self.decimal_abstract_id = self._add_decimal_abstract()

    def _next_abstract_id(self):
        ids = []
        for node in self.numbering.findall(qn("w:abstractNum")):
            val = node.get(qn("w:abstractNumId"))
            if val is not None:
                ids.append(int(val))
        return max(ids, default=0) + 1

    def _next_num_id(self):
        ids = []
        for node in self.numbering.findall(qn("w:num")):
            val = node.get(qn("w:numId"))
            if val is not None:
                ids.append(int(val))
        return max(ids, default=0) + 1

    def _add_decimal_abstract(self):
        abstract_id = self._next_abstract_id()
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))

        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)

        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "%1.")
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        lvl.append(p_pr)

        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), DOC_FONT)
        fonts.set(qn("w:hAnsi"), DOC_FONT)
        fonts.set(qn("w:eastAsia"), DOC_FONT)
        r_pr.append(fonts)
        lvl.append(r_pr)

        self.numbering.append(abstract)
        return abstract_id

    def new_decimal_num_id(self):
        num_id = self._next_num_id()
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(self.decimal_abstract_id))
        num.append(abstract_ref)
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        lvl_override.append(start_override)
        num.append(lvl_override)
        self.numbering.append(num)
        return num_id


def apply_num_pr(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_node)


DOC_FONT = "Arial Unicode MS"


def set_run_font(run, size=None, bold=None, color=None, east_asia=DOC_FONT):
    run.font.name = DOC_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, color="000000", bold=False):
    font = style.font
    font.name = DOC_FONT
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)


def set_spacing(style, before=0, after=6, line=1.25):
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_field(paragraph, instr_text):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run._r.append(instr)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    placeholder = paragraph.add_run("更新欄位後顯示")
    set_run_font(placeholder, size=10, color=MUTED)

    run_end = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)


def add_toc(paragraph):
    add_field(paragraph, r'TOC \o "1-3" \h \z \u')


def add_page_number(paragraph):
    add_field(paragraph, "PAGE")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_update_fields_on_open(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, "000000")
    set_spacing(styles["Normal"], 0, 6, 1.25)

    set_style_font(styles["Heading 1"], 16, BLUE, True)
    set_spacing(styles["Heading 1"], 18, 10, 1.25)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], 13, BLUE, True)
    set_spacing(styles["Heading 2"], 14, 7, 1.25)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], 12, DARK_BLUE, True)
    set_spacing(styles["Heading 3"], 10, 5, 1.25)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        set_style_font(styles[list_style], 11, "000000")
        set_spacing(styles[list_style], 0, 4, 1.25)
        styles[list_style].paragraph_format.left_indent = Inches(0.375)
        styles[list_style].paragraph_format.first_line_indent = Inches(-0.188)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(82)
    r = p.add_run("微型餐飲創業教學計劃")
    set_run_font(r, 26, True, INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("商業便當｜主題化 1-2 人小火鍋｜台灣特色 × 法式風格冰店")
    set_run_font(r, 14, False, DARK_BLUE)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1.65, 4.85])
    rows = [
        ("課程對象", "想小型創業但資金有限者、曾經創業未成功但仍想重新出發者"),
        ("主講者", "弘光科技大學助理教授 葉佳山"),
        ("建議時數", "18 小時，可分為 6 週；亦可濃縮為 2 日工作坊"),
        ("教材定位", "以官方資料、低成本驗證、食安法規、成本管理與產品設計為核心"),
        ("製作日期", date.today().strftime("%Y-%m-%d")),
    ]
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], TABLE_FILL)
        for i, text in enumerate((label, value)):
            para = row.cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(text)
            set_run_font(run, 10.5, i == 0, INK if i == 0 else "000000")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("給準備成為老闆的人：先小量驗證，再穩定複製。")
    set_run_font(r, 12, True, DARK_BLUE)


def setup_running_furniture(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("微型餐飲創業教學計劃")
    set_run_font(run, 9, False, MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("第 ")
    set_run_font(run, 9, False, MUTED)
    add_page_number(fp)
    run = fp.add_run(" 頁")
    set_run_font(run, 9, False, MUTED)


def normalize_inline(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return text.strip()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        parts = [normalize_inline(part) for part in lines[i].strip().strip("|").split("|")]
        if not all(re.match(r"^:?-{2,}:?$", p.strip()) for p in parts):
            rows.append(parts)
        i += 1
    if len(rows) < 1:
        return None, start
    return rows, i


def add_markdown_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    if cols == 2:
        widths = [1.875, 4.625]
    elif cols == 3:
        widths = [1.45, 2.45, 2.60]
    elif cols == 4:
        widths = [1.35, 1.65, 2.75, 0.75]
    else:
        widths = [6.5 / cols] * cols
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            if r_idx == 0:
                shade_cell(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, 9.5, r_idx == 0, INK if r_idx == 0 else "000000")
    doc.add_paragraph()


def add_source_list_item(doc, text, num_id=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    if num_id is not None:
        apply_num_pr(p, num_id)
    run = p.add_run(text)
    set_run_font(run, 10.5)


def add_content_from_markdown(doc, md_text, numbering):
    lines = md_text.splitlines()
    i = 0
    in_sources = False
    current_num_id = None
    last_was_ordered = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "---":
            last_was_ordered = False
            i += 1
            continue
        if line.startswith("# "):
            # The cover already carries the document title.
            last_was_ordered = False
            i += 1
            continue
        if line.startswith("## "):
            title = normalize_inline(line[3:])
            in_sources = title == "官方資料與課前參考"
            doc.add_paragraph(title, style="Heading 1")
            current_num_id = None
            last_was_ordered = False
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(normalize_inline(line[4:]), style="Heading 2")
            current_num_id = None
            last_was_ordered = False
            i += 1
            continue
        if line.startswith("|"):
            rows, next_i = parse_table(lines, i)
            if rows:
                add_markdown_table(doc, rows)
                last_was_ordered = False
                i = next_i
                continue
        ordered = re.match(r"^\d+\.\s+(.*)$", line)
        if ordered:
            if not last_was_ordered or current_num_id is None:
                current_num_id = numbering.new_decimal_num_id()
            text = normalize_inline(ordered.group(1))
            if in_sources and i + 1 < len(lines) and lines[i + 1].strip().startswith("http"):
                text = f"{text}：{lines[i + 1].strip()}"
                i += 1
            add_source_list_item(doc, text, current_num_id)
            last_was_ordered = True
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(normalize_inline(line[2:]))
            set_run_font(run, 10.5)
            last_was_ordered = False
            i += 1
            continue
        if line.startswith("http"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_run_font(run, 9.5, False, MUTED)
            last_was_ordered = False
            i += 1
            continue
        p = doc.add_paragraph()
        run = p.add_run(normalize_inline(line))
        set_run_font(run, 11)
        last_was_ordered = False
        i += 1


def add_toc_page(doc):
    doc.add_page_break()
    p = doc.add_paragraph("目錄", style="Heading 1")
    p.paragraph_format.space_after = Pt(8)
    toc_items = [
        "課程總目標",
        "官方資料與課前參考",
        "第一章：先把自己準備成老闆",
        "第二章：微型餐飲創業的共同基本功",
        "第三章：商業便當創業教學",
        "第四章：主題化 1-2 人小火鍋創業教學",
        "第五章：台灣特色加法式風格冰店創業教學",
        "第六章：曾經失敗者的重新出發策略",
        "第七章：課程總作業與成果發表",
        "第八章：講師結語",
        "附錄：開業前 30 日行動表",
    ]
    for item in toc_items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run, 11, False, INK if "章" in item or item.startswith("附錄") else "000000")
    doc.add_page_break()


def main():
    doc = Document()
    style_document(doc)
    numbering = NumberingManager(doc)
    set_update_fields_on_open(doc)

    doc.sections[0].different_first_page_header_footer = True
    add_cover(doc)
    add_toc_page(doc)
    setup_running_furniture(doc.sections[0])

    add_content_from_markdown(doc, SOURCE.read_text(encoding="utf-8"), numbering)

    doc.core_properties.title = "微型餐飲創業教學計劃"
    doc.core_properties.author = "弘光科技大學助理教授 葉佳山"
    doc.core_properties.subject = "商業便當、主題化小火鍋、台灣特色法式冰店微型創業教材"
    doc.core_properties.keywords = "微型創業,餐飲創業,商業便當,小火鍋,冰店,葉佳山"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
